from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE = "templates/tema_template.docx"


def _replace(para, mapping):
    full = "".join(r.text for r in para.runs)
    new = full
    for marker, value in mapping.items():
        new = new.replace(marker, value)
    if new == full or not para.runs:
        return
    para.runs[0].text = new
    for r in para.runs[1:]:
        r.text = ""


def _consolidate_runs(para):
    """Склеивает все runs параграфа в первый, сохраняя форматирование первого."""
    if not para.runs:
        return
    full = "".join(r.text for r in para.runs)
    para.runs[0].text = full
    for r in para.runs[1:]:
        r.text = ""


def _make_para_format(p_elem):
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_elem.insert(0, pPr)

    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), "both")

    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:line"), "360")
    spacing.set(qn("w:lineRule"), "auto")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")

    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:firstLine"), "709")


def _make_run(text):
    r = OxmlElement("w:r")

    rPr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:cs"), "Times New Roman")
    rPr.append(fonts)

    for tag in ("w:sz", "w:szCs"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), "28")
        rPr.append(el)

    r.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    return r


def _fill_field_26(doc, text):
    # Сначала склеиваем runs во всех параграфах чтобы метка не была разбита
    for para in doc.paragraphs:
        _consolidate_runs(para)

    target_para = None
    for para in doc.paragraphs:
        if "{{26}}" in "".join(r.text for r in para.runs):
            target_para = para
            break
    if target_para is None:
        return

    parts = [p.strip() for p in text.split("(отступ)") if p.strip()]

    parent = target_para._p.getparent()
    insert_idx = list(parent).index(target_para._p)
    parent.remove(target_para._p)

    for i, part in enumerate(parts):
        p = OxmlElement("w:p")
        _make_para_format(p)
        p.append(_make_run(part))
        parent.insert(insert_idx + i, p)


def generate(student, out_path):
    doc = Document(TEMPLATE)
    mapping = {f"{{{{{n}}}}}": v for n, v in student.items()}

    for para in doc.paragraphs:
        full = "".join(r.text for r in para.runs)
        if "{{26}}" not in full:
            _replace(para, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace(para, mapping)

    _fill_field_26(doc, student.get(26, ""))

    doc.save(out_path)
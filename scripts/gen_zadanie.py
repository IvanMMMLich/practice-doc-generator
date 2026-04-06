from docx import Document


TEMPLATE = "templates/zadanie_template.docx"


def _replace(para, mapping):
    full = "".join(r.text for r in para.runs)
    new = full
    for marker, value in mapping.items():
        new = new.replace(marker, value)
    if new == full or not para.runs:
        return
    para.runs[0].text = new
    for r in para.runs[1:]:
        r.text = ""


def _fill_stage_table(doc, s13, s14, s15):
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 3:
                continue
            label = "".join(
                r.text for r in row.cells[0].paragraphs[0].runs
            ).strip().lower()
            if "организационный" in label:
                target = s13
            elif "основной" in label:
                target = s14
            elif "заключительный" in label:
                target = s15
            else:
                continue
            cell = row.cells[2]
            for para in cell.paragraphs:
                if para.runs:
                    para.runs[0].text = target
                    for r in para.runs[1:]:
                        r.text = ""
                    break


def generate(student, out_path):
    doc = Document(TEMPLATE)
    mapping = {f"{{{{{n}}}}}": v for n, v in student.items()}

    for para in doc.paragraphs:
        _replace(para, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace(para, mapping)

    _fill_stage_table(
        doc,
        s13=student.get(13, ""),
        s14=student.get(14, ""),
        s15=student.get(15, ""),
    )

    doc.save(out_path)
from docx import Document


TEMPLATE = "templates/otzyv_template.docx"


def _replace(para, mapping):
    full = "".join(r.text for r in para.runs)
    new = full
    for marker, value in mapping.items():
        new = new.replace(marker, value)
    if new == full or not para.runs:
        return
    para.runs[0].text = new
    for r in para.runs[1:]:
        r.text = ""


def generate(student, out_path):
    doc = Document(TEMPLATE)
    mapping = {f"{{{{{n}}}}}": v for n, v in student.items()}

    for para in doc.paragraphs:
        _replace(para, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace(para, mapping)

    doc.save(out_path)